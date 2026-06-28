"""
Document Processor
Handles PDF, DOCX, XLSX, and Image files.
Returns list of {page, text} dicts for chunking.
"""
import os
import logging
from pathlib import Path
from typing import List, Dict, Any

logger = logging.getLogger(__name__)


def get_file_type(filename: str) -> str:
    ext = Path(filename).suffix.lower()
    type_map = {
        ".pdf": "pdf",
        ".docx": "docx", ".doc": "docx",
        ".xlsx": "xlsx", ".xls": "xlsx",
        ".png": "image", ".jpg": "image", ".jpeg": "image",
        ".gif": "image", ".bmp": "image", ".tiff": "image", ".webp": "image",
    }
    return type_map.get(ext, "unknown")


async def process_document(file_path: str, filename: str) -> List[Dict[str, Any]]:
    """
    Returns a list of page dicts:
    [{"page": 1, "text": "...", "metadata": {...}}, ...]
    """
    file_type = get_file_type(filename)
    logger.info(f"Processing {file_type}: {filename}")

    if file_type == "pdf":
        return await _process_pdf(file_path, filename)
    elif file_type == "docx":
        return await _process_docx(file_path, filename)
    elif file_type == "xlsx":
        return await _process_xlsx(file_path, filename)
    elif file_type == "image":
        return await _process_image(file_path, filename)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")


async def _process_pdf(file_path: str, filename: str) -> List[Dict]:
    import fitz  # PyMuPDF
    pages = []
    doc = fitz.open(file_path)
    for page_num in range(len(doc)):
        page = doc[page_num]
        text = page.get_text("text").strip()
        if text:
            pages.append({
                "page": page_num + 1,
                "text": text,
                "metadata": {
                    "filename": filename,
                    "page": page_num + 1,
                    "total_pages": len(doc),
                    "source": filename,
                }
            })
    doc.close()
    return pages


async def _process_docx(file_path: str, filename: str) -> List[Dict]:
    from docx import Document
    doc = Document(file_path)
    # Group paragraphs into "pages" of ~500 words
    all_text = []
    for para in doc.paragraphs:
        if para.text.strip():
            all_text.append(para.text.strip())
    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                all_text.append(row_text)

    # Split into pseudo-pages of ~500 words
    pages = []
    chunk, word_count, page_num = [], 0, 1
    for line in all_text:
        words = len(line.split())
        chunk.append(line)
        word_count += words
        if word_count >= 500:
            pages.append({
                "page": page_num,
                "text": "\n".join(chunk),
                "metadata": {"filename": filename, "page": page_num, "source": filename}
            })
            chunk, word_count, page_num = [], 0, page_num + 1
    if chunk:
        pages.append({
            "page": page_num,
            "text": "\n".join(chunk),
            "metadata": {"filename": filename, "page": page_num, "source": filename}
        })
    return pages


async def _process_xlsx(file_path: str, filename: str) -> List[Dict]:
    import openpyxl
    wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
    pages = []
    for sheet_idx, sheet_name in enumerate(wb.sheetnames):
        ws = wb[sheet_name]
        rows_text = []
        for row in ws.iter_rows(values_only=True):
            row_vals = [str(cell) for cell in row if cell is not None and str(cell).strip()]
            if row_vals:
                rows_text.append(" | ".join(row_vals))
        if rows_text:
            pages.append({
                "page": sheet_idx + 1,
                "text": f"[Sheet: {sheet_name}]\n" + "\n".join(rows_text),
                "metadata": {
                    "filename": filename,
                    "page": sheet_idx + 1,
                    "sheet": sheet_name,
                    "source": filename,
                }
            })
    wb.close()
    return pages


async def _process_image(file_path: str, filename: str) -> List[Dict]:
    try:
        import easyocr
        import numpy as np
        from PIL import Image

        img = Image.open(file_path).convert("RGB")
        img_array = np.array(img)
        reader = easyocr.Reader(["en"], gpu=True, verbose=False)
        results = reader.readtext(img_array, detail=0, paragraph=True)
        text = "\n".join(results)
        if not text.strip():
            text = "[Image: No readable text detected]"
    except Exception as e:
        logger.warning(f"EasyOCR failed ({e}), using placeholder text")
        text = f"[Image file: {filename}. OCR processing failed.]"

    return [{
        "page": 1,
        "text": text,
        "metadata": {"filename": filename, "page": 1, "source": filename}
    }]
